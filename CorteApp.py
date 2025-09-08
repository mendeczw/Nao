import os
from datetime import datetime
import pandas as pd
from tkinter import Tk, Toplevel, Label, Button, filedialog, messagebox, StringVar
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------- Configuración ----------------
APP_NAME = "Generador de Análisis Ejecutivo"
CREDITO = "Realizado por William Méndez – mendezcw@yahoo.es"
RUTA_INICIAL = r"C:\Users\mende\OneDrive\Documents\ANALISIS_DAV"
NOMBRE_HOJA_POR_DEFECTO = None  # None = primera hoja

# Columnas esperadas (variantes toleradas -> nombre estándar)
ALIASES = {
    "super": "SUPER",
    "usuario": "Usuario",
    "registros recorrido": "Registros Recorrido",
    "registros recorridos": "Registros Recorrido",
    "spin rate": "Spin Rate",
    "contactados": "Contactados",
    "contacto efectivo": "Contacto Efectivo",
    "contacto no valido": "Contacto No Valido",
    # errores comunes en origen
    "no contatados": "No Contactados",
    "no contactados": "No Contactados",
    "venta": "Venta",
    "no aplica": "No Aplica",
    "% contactabilidad": "% Contactabilidad",
    "%c/efectivo": "%C/Efectivo",
    "conversion": "Conversión",
    "conversión": "Conversión",
    "penetracion": "Penetración",
    "penetración": "Penetración",
}

COLS_MINIMAS = [
    "Usuario", "Registros Recorrido", "Contactados", "Contacto Efectivo",
    "Contacto No Valido", "No Contactados", "Venta", "No Aplica"
]

# ---------------- Lógica de análisis ----------------
def normalizar_cols(cols):
    out = []
    for c in cols:
        s = str(c).strip().replace("\n", " ")
        s = (s
             .replace("á", "a").replace("é","e").replace("í","i")
             .replace("ó","o").replace("ú","u").replace("ñ","n"))
        s_low = s.lower()
        out.append(ALIASES.get(s_low, s))  # mapear si coincide en minúsculas
    return out

def cargar_excel(ruta, sheet_name=NOMBRE_HOJA_POR_DEFECTO):
    xl = pd.ExcelFile(ruta)
    hoja = sheet_name if (sheet_name and sheet_name in xl.sheet_names) else (sheet_name or xl.sheet_names[0])
    df = xl.parse(hoja)
    df.columns = normalizar_cols(df.columns)
    # quitar fila de totales si existe
    if "Usuario" in df.columns:
        df = df[~df["Usuario"].astype(str).str.lower().str.contains("totales", na=False)]
    return df

def asegurar_columnas(df):
    # coerción inicial a numérico para columnas clave si existen
    for c in ["Registros Recorrido", "Contactados"]:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0)

    # si falta 'No Contactados' pero hay datos para inferirlo, lo calculo
    if "No Contactados" not in df.columns and \
       "Registros Recorrido" in df.columns and "Contactados" in df.columns:
        df["No Contactados"] = df["Registros Recorrido"] - df["Contactados"]

    faltan = [c for c in COLS_MINIMAS if c not in df.columns]
    if faltan:
        raise ValueError(f"Faltan columnas obligatorias: {faltan}\n"
                         f"Columnas presentes: {list(df.columns)}")

def kpis_generales(df):
    asegurar_columnas(df)

    # Coerción a numérico total
    for c in ["Registros Recorrido","Contactados","Contacto Efectivo",
              "Contacto No Valido","No Contactados","Venta","No Aplica"]:
        df[c] = pd.to_numeric(df[c], errors="coerce").fillna(0)

    tot_recorridos = int(df["Registros Recorrido"].sum())
    tot_contactados = int(df["Contactados"].sum())
    tot_efectivos = int(df["Contacto Efectivo"].sum())
    tot_no_validos = int(df["Contacto No Valido"].sum())
    tot_no_contactados = int(df["No Contactados"].sum())
    tot_ventas = int(df["Venta"].sum())
    tot_no_aplica = int(df["No Aplica"].sum())

    contactabilidad = (tot_contactados / tot_recorridos) if tot_recorridos else 0.0
    efectividad = (tot_efectivos / tot_contactados) if tot_contactados else 0.0
    conversion = (tot_ventas / tot_efectivos) if tot_efectivos else 0.0
    penetracion = (tot_ventas / tot_recorridos) if tot_recorridos else 0.0

    return {
        "tot_recorridos": tot_recorridos,
        "tot_contactados": tot_contactados,
        "tot_efectivos": tot_efectivos,
        "tot_no_validos": tot_no_validos,
        "tot_no_contactados": tot_no_contactados,
        "tot_ventas": tot_ventas,
        "tot_no_aplica": tot_no_aplica,
        "contactabilidad": contactabilidad,
        "efectividad": efectividad,
        "conversion": conversion,
        "penetracion": penetracion,
    }

def diagnostico_texto(kpi):
    msgs = []
    if kpi["contactabilidad"] < 0.35:
        msgs.append(f"La contactabilidad es baja ({kpi['contactabilidad']:.0%}). Urge mejorar horarios, marcación multicanal y calidad de bases.")
    else:
        msgs.append(f"Buena contactabilidad ({kpi['contactabilidad']:.0%}). Mantener estrategia y ampliar ventanas horarias.")
    if kpi["efectividad"] >= 0.85:
        msgs.append(f"La efectividad sobre contactados es sólida ({kpi['efectividad']:.0%}). El discurso funciona.")
    else:
        msgs.append(f"La efectividad es mejorable ({kpi['efectividad']:.0%}). Revisar objeciones y guion.")
    if kpi["conversion"] < 0.10:
        msgs.append(f"La conversión sobre contactos efectivos es baja ({kpi['conversion']:.1%}). Ajustar oferta, cross-sell y cierres.")
    else:
        msgs.append(f"Conversión saludable ({kpi['conversion']:.1%}). Replicar mejores prácticas.")
    msgs.append(f"Penetración global: {kpi['penetracion']:.1%} sobre la base recorrida.")
    return msgs

def top_agentes(df, n=5):
    base = df.copy()
    for c in ["Registros Recorrido","Contactados","Contacto Efectivo","Venta"]:
        if c in base.columns:
            base[c] = pd.to_numeric(base[c], errors="coerce").fillna(0)

    base["%Efectividad"] = base["Contacto Efectivo"] / base["Contactados"].replace(0, pd.NA)
    base["%Conversión"] = base["Venta"] / base["Contacto Efectivo"].replace(0, pd.NA)

    # Orden principal por ventas (desc), desempate por %Efectividad (desc)
    orden = base.sort_values(["Venta","%Efectividad"], ascending=[False, False])

    cols_out = ["Usuario","Registros Recorrido","Contactados","Contacto Efectivo",
                "Venta","%Efectividad","%Conversión"]
    cols_out = [c for c in cols_out if c in orden.columns]
    return orden[cols_out].head(n)

def comentario_top(top):
    comentarios = []
    if top.empty:
        return ["No hay datos para análisis comparativo del Top 5."]

    # líder en ventas (top ya viene ordenado por ventas desc)
    lider_ventas = top.iloc[0]

    # mejor efectividad
    idx_eff = top["%Efectividad"].idxmax()
    mejor_eff = top.loc[idx_eff]

    comentarios.append(
        f"{lider_ventas['Usuario']} lidera en ventas con {int(lider_ventas['Venta'])} cierres."
    )
    comentarios.append(
        f"{mejor_eff['Usuario']} presenta la mayor efectividad ({mejor_eff['%Efectividad']:.0%}) en contactos."
    )

    if lider_ventas["Usuario"] == mejor_eff["Usuario"]:
        comentarios.append(
            f"Se confirma un desempeño integral: {lider_ventas['Usuario']} lidera tanto en volumen como en eficiencia."
        )
    else:
        comentarios.append(
            f"Hallazgo: {lider_ventas['Usuario']} domina el volumen de cierres, "
            f"mientras {mejor_eff['Usuario']} destaca por la eficiencia. Conviene compartir prácticas de ambos."
        )
    return comentarios

# --------------- Word (solo texto + top + comentario) ---------------
def agregar_parrafo(doc, texto, bold=False, size=11, align_left=True):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = bold
    run.font.size = Pt(size)
    if not align_left:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def generar_word(df, kpi, ruta_excel):
    nombre_archivo = os.path.splitext(os.path.basename(ruta_excel))[0]
    ahora = datetime.now().strftime("%Y-%m-%d %H.%M")
    out_dir = os.path.dirname(ruta_excel)
    salida = os.path.join(out_dir, f"Reporte_Corte_{nombre_archivo}_{ahora}.docx")

    doc = Document()
    agregar_parrafo(doc, "Análisis General del Corte", bold=True, size=16, align_left=False)
    agregar_parrafo(doc, f"Archivo origen: {os.path.basename(ruta_excel)}", size=9)
    agregar_parrafo(doc, f"Fecha de generación: {datetime.now():%d/%m/%Y %H:%M}", size=9)
    agregar_parrafo(doc, CREDITO, size=9)

    doc.add_paragraph()

    # Resumen ejecutivo (solo texto)
    agregar_parrafo(doc, "Resumen Ejecutivo", bold=True, size=13)
    doc.add_paragraph(
        f"Registros recorridos: {kpi['tot_recorridos']:,}".replace(",", "."))
    doc.add_paragraph(
        f"Contactados: {kpi['tot_contactados']:,} | Contacto Efectivo: {kpi['tot_efectivos']:,} | No válidos: {kpi['tot_no_validos']:,}".replace(",", "."))
    doc.add_paragraph(
        f"No contactados: {kpi['tot_no_contactados']:,} | Ventas: {kpi['tot_ventas']:,} | No aplica: {kpi['tot_no_aplica']:,}".replace(",", "."))

    doc.add_paragraph()

    # Diagnóstico (viñetas)
    agregar_parrafo(doc, "Diagnóstico", bold=True, size=13)
    for linea in diagnostico_texto(kpi):
        doc.add_paragraph(f"• {linea}")

    doc.add_paragraph()

    # Top 5 Agentes (listado simple)
    agregar_parrafo(doc, "Top 5 Agentes", bold=True, size=13)
    top = top_agentes(df, n=5)
    for _, row in top.iterrows():
        eff = row["%Efectividad"]
        conv = row["%Conversión"]
        eff_txt = f"{float(eff):.0%}" if pd.notna(eff) else "N/D"
        conv_txt = f"{float(conv):.0%}" if pd.notna(conv) else "N/D"

        doc.add_paragraph(
            f"{row['Usuario']}: "
            f"Recorridos {int(row['Registros Recorrido'])}, "
            f"Contactados {int(row['Contactados'])}, "
            f"Efectivos {int(row['Contacto Efectivo'])}, "
            f"Ventas {int(row['Venta'])}, "
            f"Efectividad {eff_txt}, "
            f"Conversión {conv_txt}"
        )

    # Comentario sobre el Top 5
    doc.add_paragraph()
    agregar_parrafo(doc, "Comentario sobre el Top 5", bold=True, size=13)
    for linea in comentario_top(top):
        doc.add_paragraph(f"• {linea}")

    doc.save(salida)
    return salida

# ----------------- GUI Tkinter -----------------
class CorteApp(Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_NAME)
        self.geometry("620x220")
        self.resizable(False, False)

        self.ruta_excel = None
        self.status = StringVar(value="Seleccione un archivo y genere el reporte.")
        self.lbl_title = Label(self, text=APP_NAME, font=("Segoe UI", 14, "bold"))
        self.lbl_title.pack(pady=(12, 4))

        self.lbl_credito = Label(self, text=CREDITO, font=("Segoe UI", 9))
        self.lbl_credito.pack(pady=(0, 8))

        self.lbl_fecha = Label(self, text=f"Fecha: {datetime.now():%d/%m/%Y}", font=("Segoe UI", 9))
        self.lbl_fecha.pack(pady=(0, 8))

        self.btn_archivo = Button(self, text="Seleccionar archivo Excel...", width=40, command=self.seleccionar_archivo)
        self.btn_archivo.pack(pady=4)

        self.btn_generar = Button(self, text="Generar Reporte Word", width=40, command=self.generar_reporte)
        self.btn_generar.pack(pady=4)

        self.lbl_status = Label(self, textvariable=self.status, font=("Segoe UI", 9))
        self.lbl_status.pack(pady=(8, 6))

    def seleccionar_archivo(self):
        ruta = filedialog.askopenfilename(
            initialdir=RUTA_INICIAL,
            title="Seleccione el archivo de corte",
            filetypes=[("Excel", "*.xlsx *.xlsm *.xls *.xlsb"), ("Todos", "*.*")]
        )
        if ruta:
            self.ruta_excel = ruta
            self.status.set(f"Archivo seleccionado: {os.path.basename(ruta)}")

    def generar_reporte(self):
        if not self.ruta_excel:
            messagebox.showwarning("Atención", "Primero seleccione un archivo Excel.")
            return
        try:
            # Detectar hoja si hay varias (tomar la primera)
            df = cargar_excel(self.ruta_excel, sheet_name=NOMBRE_HOJA_POR_DEFECTO)
            kpi = kpis_generales(df)
            salida = generar_word(df, kpi, self.ruta_excel)
            self.status.set(f"Reporte generado: {salida}")
            messagebox.showinfo("Proceso finalizado", f"Reporte generado correctamente.\n\nRuta:\n{salida}")
        except Exception as e:
            self.status.set("Ocurrió un error al generar el reporte.")
            messagebox.showerror("Error", str(e))

def main():
    app = CorteApp()
    app.mainloop()

if __name__ == "__main__":
    main()
